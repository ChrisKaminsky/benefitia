from docx import Document
from bs4 import BeautifulSoup


def convert_docx_to_html(docx_file):
    # Open the .docx file
    doc = Document(docx_file)

    # Create a BeautifulSoup object for HTML parsing
    soup = BeautifulSoup(features="html.parser")

    # Process paragraphs and lists (bullet points)
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('List'):
            ul_tag = soup.new_tag('ul')
            li_tag = soup.new_tag('li')
            li_tag.string = paragraph.text
            ul_tag.append(li_tag)
            soup.append(ul_tag)
        else:
            p_tag = soup.new_tag('p')
            p_tag.string = paragraph.text
            soup.append(p_tag)

    # Process tables
    for table in doc.tables:
        table_tag = soup.new_tag('table')

        for row in table.rows:
            tr_tag = soup.new_tag('tr')

            for cell in row.cells:
                td_tag = soup.new_tag('td')
                td_tag.string = cell.text
                tr_tag.append(td_tag)

            table_tag.append(tr_tag)

        soup.append(table_tag)

    # Convert BeautifulSoup object to HTML string
    html_output = soup.prettify()

    return html_output


# Usage example
docx_file_path = 'input.docx'
html_output = convert_docx_to_html(docx_file_path)

# Save HTML output to a file
with open('output.html', 'w', encoding='utf-8') as file:
    file.write(html_output)
